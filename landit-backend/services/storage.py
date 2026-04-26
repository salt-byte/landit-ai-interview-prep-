"""
File storage service. Uses local filesystem by default.
Can be swapped to S3 by replacing upload_file/get_file_path.
"""
import asyncio
import re
import uuid
import aiofiles
from pathlib import Path
from fastapi import UploadFile
from config import settings


async def upload_file(file: UploadFile, subfolder: str = "") -> tuple[str, int]:
    """
    Save uploaded file to local storage.
    Returns (file_path, file_size_bytes).
    """
    upload_dir = settings.upload_path
    if subfolder:
        upload_dir = upload_dir / subfolder
        upload_dir.mkdir(parents=True, exist_ok=True)

    # Generate unique filename preserving extension
    ext = Path(file.filename or "file").suffix
    unique_name = f"{uuid.uuid4().hex}{ext}"
    dest = upload_dir / unique_name

    content = await file.read()
    async with aiofiles.open(dest, "wb") as f:
        await f.write(content)

    return str(dest), len(content)


def get_file_path(stored_path: str) -> Path:
    return Path(stored_path)


async def read_text_file(file_path: str) -> str:
    """Read a text file (for resume parsing, etc.)"""
    path = Path(file_path)
    if not path.exists():
        return ""

    if path.suffix.lower() == ".pdf":
        try:
            return await asyncio.to_thread(_read_pdf_text, path)
        except Exception:
            return ""

    if path.suffix.lower() == ".docx":
        try:
            return await asyncio.to_thread(_read_docx_text, path)
        except Exception:
            return ""

    async with aiofiles.open(path, "r", errors="ignore") as f:
        return await f.read()


# Resumes use various decorative characters as bullet points. pypdf extracts
# them as raw codepoints which then leak into the LLM output as garbage.
# We strip them entirely — the LLM will add its own bullets when emitting
# the description field.
#
# Covers:
#   Math Operators        U+2200–U+22FF  (∀ … ⋿, includes ≡)
#   Geometric Shapes      U+25A0–U+25FF  (■ ● ◆ ▪ ▶)
#   Misc Symbols/Dingbats U+2600–U+27BF  (★ ➤ ✓)
#   I Ching trigrams      U+2630–U+2637  (☰)
#   Private Use Area      U+E000–U+F8FF  (Wingdings/Symbol custom glyphs)
_DECORATIVE_CHARS = re.compile(
    "["
    "∀-⋿"
    "■-◿"
    "☀-➿"
    "☰-☷"
    "-"
    "]"
)


def _clean_pdf_text(raw: str) -> str:
    cleaned = _DECORATIVE_CHARS.sub("", raw)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\s+\n", "\n", cleaned)
    cleaned = re.sub(r"\n\s+", "\n", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _read_pdf_text(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    text = "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    return _clean_pdf_text(text)


def _read_docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs).strip()


async def delete_file(file_path: str) -> bool:
    path = Path(file_path)
    if path.exists():
        path.unlink()
        return True
    return False


def detect_file_type(filename: str) -> str:
    """Heuristic file type detection by name."""
    name = filename.lower()
    if any(x in name for x in ["resume", "cv"]):
        return "Resume"
    if any(x in name for x in ["portfolio", "case", "project"]):
        return "Portfolio"
    if any(x in name for x in ["sql", "code", "sample", "script", "py", "js"]):
        return "Work Sample"
    if any(x in name for x in ["jd", "job", "description", "posting"]):
        return "Job Description"
    return "Notes"
